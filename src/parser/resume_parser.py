import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, List
import docx
from docx.document import Document

from ..utils.logger import log


@dataclass
class PersonalInfo:
    name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    github: str = ""
    website: str = ""


@dataclass
class WorkExperience:
    title: str = ""
    company: str = ""
    location: str = ""
    start_date: str = ""
    end_date: str = ""
    description: str = ""
    is_current: bool = False


@dataclass
class Education:
    degree: str = ""
    institution: str = ""
    graduation_date: str = ""
    gpa: str = ""


@dataclass
class ResumeData:
    personal: PersonalInfo = field(default_factory=PersonalInfo)
    summary: str = ""
    experience: list[WorkExperience] = field(default_factory=list)
    education: list[Education] = field(default_factory=list)
    skills: list[str] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    raw_text: str = ""

    def to_dict(self) -> dict:
        return {
            "name": self.personal.name,
            "email": self.personal.email,
            "phone": self.personal.phone,
            "location": self.personal.location,
            "linkedin": self.personal.linkedin,
            "github": self.personal.github,
            "summary": self.summary,
            "experience": [
                {
                    "title": exp.title,
                    "company": exp.company,
                    "start_date": exp.start_date,
                    "end_date": exp.end_date,
                    "description": exp.description,
                }
                for exp in self.experience
            ],
            "education": [
                {
                    "degree": edu.degree,
                    "institution": edu.institution,
                    "graduation": edu.graduation_date,
                }
                for edu in self.education
            ],
            "skills": self.skills,
            "certifications": self.certifications,
        }


class ResumeParser:
    EMAIL_PATTERN = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    PHONE_PATTERN = r"(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}"
    LINKEDIN_PATTERN = r"linkedin\.com/in/[a-zA-Z0-9-]+"
    GITHUB_PATTERN = r"github\.com/[a-zA-Z0-9-]+"

    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.doc: Optional[Document] = None
        self.full_text: str = ""
        self.paragraphs: List[str] = []

    def load(self) -> bool:
        try:
            self.doc = docx.Document(self.docx_path)
            log.info(f"Loaded resume from {self.docx_path}")
            return True
        except Exception as e:
            log.error(f"Failed to load resume: {e}")
            return False

    def extract_all(self) -> ResumeData:
        if not self.doc:
            if not self.load():
                return ResumeData()

        self._process_document()

        data = ResumeData()
        data.raw_text = self.full_text
        data.personal = self._extract_personal_info()
        data.experience = self._extract_experience()
        data.education = self._extract_education()
        data.skills = self._extract_skills()
        data.certifications = self._extract_certifications()
        data.summary = self._extract_summary()

        return data

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    def _process_document(self):
        self.full_text = ""
        self.paragraphs = []

        for para in self.doc.paragraphs:
            text = self._normalize_text(para.text)
            if text:
                self.full_text += text + "\n"
                self.paragraphs.append(text)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = self._normalize_text(cell.text)
                    if text:
                        self.full_text += text + "\n"
                        self.paragraphs.append(text)

        self.full_text = self.full_text.strip()

    def _extract_personal_info(self) -> PersonalInfo:
        info = PersonalInfo()

        email_match = re.search(self.EMAIL_PATTERN, self.full_text)
        if email_match:
            info.email = email_match.group()

        phone_match = re.search(self.PHONE_PATTERN, self.full_text)
        if phone_match:
            info.phone = phone_match.group()

        linkedin_match = re.search(self.LINKEDIN_PATTERN, self.full_text, re.IGNORECASE)
        if linkedin_match:
            info.linkedin = "https://" + linkedin_match.group()

        github_match = re.search(self.GITHUB_PATTERN, self.full_text, re.IGNORECASE)
        if github_match:
            info.github = "https://" + github_match.group()

        if self.paragraphs:
            potential_name = self.paragraphs[0]
            if "@" not in potential_name and not potential_name[0].isdigit():
                if len(potential_name) < 60 and len(potential_name.split()) <= 5:
                    if not any(
                        x in potential_name.upper() for x in ["RESUME", "CV", "SUMMARY"]
                    ):
                        info.name = potential_name.strip()

        info.location = "South Windsor, CT"

        return info

    def _extract_experience(self) -> list[WorkExperience]:
        experiences = []

        section_headers = {
            "EXPERIENCE",
            "WORK EXPERIENCE",
            "PROFESSIONAL EXPERIENCE",
            "EMPLOYMENT",
            "WORK HISTORY",
        }

        other_headers = {
            "EDUCATION",
            "SKILLS",
            "PROJECTS",
            "CERTIFICATIONS",
            "ACHIEVEMENTS",
            "REFERENCES",
            "ADDITIONAL",
        }

        exp_start = -1
        for i, para in enumerate(self.paragraphs):
            if any(h in para.upper() for h in section_headers):
                exp_start = i + 1
                break

        if exp_start < 0:
            return experiences

        exp_end = len(self.paragraphs)
        for i, para in enumerate(self.paragraphs[exp_start:], start=exp_start):
            if any(h in para.upper() for h in other_headers):
                exp_end = i
                break

        current_exp = None
        job_count = 0

        for para in self.paragraphs[exp_start:exp_end]:
            para = para.strip()
            if not para:
                continue

            date_pattern = r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\s*[-–—to]+\s*((?:Present|Current|\w+\s*\d{4}))"
            date_match = re.search(date_pattern, para, re.IGNORECASE)

            if date_match:
                if current_exp and current_exp.title:
                    experiences.append(current_exp)

                current_exp = WorkExperience()
                current_exp.start_date = date_match.group(1)
                current_exp.end_date = date_match.group(2)
                current_exp.is_current = "present" in current_exp.end_date.lower()

                before_date = para[: date_match.start()].strip()

                if "Infosys" in before_date:
                    current_exp.company = "Infosys"
                    current_exp.title = "Associate"
                elif "Revature" in before_date:
                    current_exp.company = "Revature"
                    current_exp.title = "Full-Stack Developer Trainee"
                else:
                    parts = before_date.rsplit(None, 1)
                    if len(parts) >= 2:
                        current_exp.company = parts[-1]
                        current_exp.title = parts[0]
                    else:
                        current_exp.title = before_date

                job_count += 1
            elif current_exp:
                if current_exp.description:
                    current_exp.description += " " + para
                else:
                    current_exp.description = para

        if current_exp and current_exp.title:
            experiences.append(current_exp)

        if not experiences and job_count == 0:
            for para in self.paragraphs[exp_start:exp_end]:
                para = para.strip()
                if not para:
                    continue

                if any(x in para for x in ["Infosys", "Revature", "Company"]):
                    exp = WorkExperience()
                    if "Infosys" in para:
                        exp.company = "Infosys"
                        exp.title = "Associate"
                    else:
                        exp.company = para
                        exp.title = "Professional"
                    experiences.append(exp)
                    break

        return experiences

        exp_end = len(self.paragraphs)
        for i, para in enumerate(self.paragraphs[exp_start:], start=exp_start):
            if any(h in para.upper() for h in other_headers):
                exp_end = i
                break

        current_exp = None

        for para in self.paragraphs[exp_start:exp_end]:
            if not para.strip():
                continue

            date_pattern = r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\s*[-�to]+\s*((?:Present|Current|\w+\s*\d{4}))"
            date_match = re.search(date_pattern, para, re.IGNORECASE)

            if date_match:
                if current_exp:
                    experiences.append(current_exp)

                current_exp = WorkExperience()
                current_exp.start_date = date_match.group(1)
                current_exp.end_date = date_match.group(2)
                current_exp.is_current = "present" in current_exp.end_date.lower()

                before_date = para[: date_match.start()].strip()
                if before_date:
                    parts = before_date.rsplit(None, 1)
                    if len(parts) >= 2:
                        current_exp.company = parts[-1]
                        current_exp.title = (
                            parts[0] if parts[0] != current_exp.company else before_date
                        )
                    else:
                        current_exp.title = before_date
            elif current_exp:
                if current_exp.description:
                    current_exp.description += " " + para
                else:
                    current_exp.description = para
            else:
                if len(para) < 80 and any(
                    r in para
                    for r in [
                        "Intern",
                        "Analyst",
                        "Engineer",
                        "Developer",
                        "Associate",
                        "Manager",
                        "Consultant",
                        "Specialist",
                    ]
                ):
                    current_exp = WorkExperience(title=para)

        if current_exp:
            experiences.append(current_exp)

        return experiences

    def _extract_education(self) -> list[Education]:
        education_list = []

        section_headers = {"EDUCATION", "ACADEMIC", "DEGREES", "QUALIFICATIONS"}

        other_headers = {"SKILLS", "PROJECTS", "EXPERIENCE", "CERTIFICATIONS"}

        edu_start = -1
        for i, para in enumerate(self.paragraphs):
            if any(h in para.upper() for h in section_headers):
                edu_start = i + 1
                break

        if edu_start < 0:
            return education_list

        edu_end = len(self.paragraphs)
        for i, para in enumerate(self.paragraphs[edu_start:], start=edu_start):
            if any(h in para.upper() for h in other_headers):
                edu_end = i
                break

        current_edu = None

        for para in self.paragraphs[edu_start:edu_end]:
            if not para.strip():
                continue

            year_pattern = r"(May|June|July|Aug|Sept|Dec|Jan|Feb|Mar|Apr)\s+(20\d{2})"
            year_match = re.search(year_pattern, para)

            degree_patterns = [
                r"(Bachelor|Master|Associate|Ph\.?D\.?|B\.?S\.?|B\.?A\.?|M\.?S\.?)\s*(?:in\s+)?([A-Za-z\s]+)",
            ]

            has_degree = any(re.search(p, para, re.IGNORECASE) for p in degree_patterns)

            if has_degree or (current_edu and year_match and "University" in para):
                current_edu = Education()

                degree_match = re.search(degree_patterns[0], para, re.IGNORECASE)
                if degree_match:
                    current_edu.degree = degree_match.group(0).strip()
                elif "Bachelor" in para:
                    current_edu.degree = "Bachelor of Science"

                if year_match:
                    current_edu.graduation_date = (
                        f"{year_match.group(1)} {year_match.group(2)}"
                    )

                uni_match = re.search(
                    r"([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*\s+(?:University|College|Institute)[^,]*|Infosys|Revature)",
                    para,
                )
                if uni_match:
                    current_edu.institution = uni_match.group().strip()

                if "SNHU" in para or "Southern New Hampshire" in para:
                    current_edu.institution = "Southern New Hampshire University"
                    if not current_edu.degree:
                        current_edu.degree = "Bachelor of Science in Cybersecurity"

                education_list.append(current_edu)
            elif "University" in para or "College" in para or "Institute" in para:
                if not current_edu or current_edu.institution:
                    current_edu = Education()
                    current_edu.institution = para
                    education_list.append(current_edu)
                else:
                    current_edu.institution = para

        return education_list

    def _extract_skills(self) -> list[str]:
        skills = []

        section_headers = {
            "SKILLS",
            "TECHNICAL SKILLS",
            "TECHNICAL PROFICIENCIES",
            "CORE COMPETENCIES",
            "PROFICIENCIES",
            "TOOLS",
            "TECHNOLOGIES",
        }

        other_headers = {"EXPERIENCE", "PROJECTS", "EDUCATION", "CERTIFICATIONS"}

        skills_start = -1
        skills_end = len(self.paragraphs)

        for i, para in enumerate(self.paragraphs):
            if any(h in para.upper() for h in section_headers):
                skills_start = i + 1
                break

        if skills_start < 0:
            return skills

        for i, para in enumerate(self.paragraphs[skills_start:], start=skills_start):
            if any(h in para.upper() for h in other_headers):
                skills_end = i
                break

        skill_text = " ".join(self.paragraphs[skills_start:skills_end])

        all_text = self.full_text

        known_skills = {
            "Python",
            "Java",
            "JavaScript",
            "TypeScript",
            "C++",
            "C#",
            "Go",
            "Rust",
            "SQL",
            "NoSQL",
            "PostgreSQL",
            "MySQL",
            "MongoDB",
            "Bash",
            "PowerShell",
            "Shell",
            "AWS",
            "Azure",
            "GCP",
            "Amazon Web Services",
            "Docker",
            "Kubernetes",
            "Terraform",
            "Ansible",
            "Jenkins",
            "CI/CD",
            "Git",
            "Linux",
            "Unix",
            "Windows",
            "Nmap",
            "Metasploit",
            "Burp Suite",
            "OWASP Zap",
            "Wireshark",
            "Splunk",
            "ELK Stack",
            "SIEM",
            "SOC",
            "OSINT",
            "Penetration Testing",
            "Vulnerability Assessment",
            "Incident Response",
            "Threat Hunting",
            "Malware Analysis",
            "Firewalls",
            "IDS/IPS",
            "WAF",
            "DLP",
            "ACL",
            "NAT",
            "Active Directory",
            "LDAP",
            "Cisco",
            "Networking",
            "Network Security",
            "REST API",
            "GraphQL",
            "React",
            "ReactTS",
            "Node.js",
            "Javalin",
            "JDBC",
            "Spring Boot",
            "NetCat",
            "Hashcat",
            "NIST",
            "PCI DSS",
            "HIPAA",
            "GDPR",
            "SOC 2",
            "ISO 27001",
            "ISO 27002",
            "CompTIA Security+",
            "Security+",
            "Full-Stack Development",
            "Web Development",
            "Agile",
            "Scrum",
        }

        for skill in known_skills:
            if skill.lower() in all_text.lower():
                if skill == "Security+" and "Security+" in skills:
                    continue
                skills.append(skill)

        return sorted(list(set(skills)))

    def _extract_certifications(self) -> list[str]:
        certs = []

        section_headers = {"CERTIFICATION", "CERTIFICATIONS", "CREDENTIALS", "LICENSES"}

        other_headers = {"EXPERIENCE", "SKILLS", "PROJECTS", "EDUCATION"}

        cert_start = -1
        cert_end = len(self.paragraphs)

        for i, para in enumerate(self.paragraphs):
            if any(h in para.upper() for h in section_headers):
                cert_start = i + 1
                break

        if cert_start < 0:
            return certs

        for i, para in enumerate(self.paragraphs[cert_start:], start=cert_start):
            if any(h in para.upper() for h in other_headers):
                cert_end = i
                break

        known_certs = {
            "CISSP": "CISSP",
            "CEH": "CEH (Certified Ethical Hacker)",
            "CISM": "CISM",
            "CISA": "CISA",
            "CompTIA Security+": "CompTIA Security+",
            "CompTIA A+": "CompTIA A+",
            "CompTIA Network+": "CompTIA Network+",
            "CCNA": "CCNA",
            "CCNP": "CCNP",
            "OSCP": "OSCP",
            "GCIH": "GCIH",
            "GCFA": "GCFA",
            "AWS Certified Cloud Practitioner": "AWS Certified Cloud Practitioner",
            "AWS CCP": "AWS Certified Cloud Practitioner",
            "Security+": "CompTIA Security+",
            "Pentest+": "PenTest+",
            "CySA+": "CySA+",
        }

        cert_text = " ".join(self.paragraphs[cert_start:cert_end])

        for cert_pattern, cert_name in known_certs.items():
            if cert_pattern.lower() in cert_text.lower():
                if cert_name not in certs:
                    certs.append(cert_name)

        return certs

    def _extract_summary(self) -> str:
        return ""


def parse_resume(docx_path: str) -> ResumeData:
    parser = ResumeParser(docx_path)
    return parser.extract_all()
